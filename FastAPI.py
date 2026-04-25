from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from retriever import HybridRetriever
from pydantic import BaseModel
from typing import Literal
import ollama
import docx  # python-docx
import PyPDF2
from odfdo import Document as OdfDocument
from io import BytesIO
import os
import json
import re

def decode_file_contents(contents: bytes) -> str:
    encodings_to_try = [
        'utf-8',
        'cp1251',  # Windows-1251
        'koi8-r',  # Русская KOI8
        'iso-8859-5',  # Latin/Cyrillic
        'windows-1252',  # Западная Европа
        'ascii',  # Только ASCII
    ]

    for encoding in encodings_to_try:
        try:
            return contents.decode(encoding)
        except UnicodeDecodeError:
            continue

    # Если ни одна кодировка не подошла, ищем читаемые фрагменты
    try:
        # Пытаемся декодировать как UTF-8, заменяя проблемные символы
        return contents.decode('utf-8', errors='replace')
    except:
        raise HTTPException(
            status_code=400,
            detail="Unable to decode file: unsupported encoding. Please provide a plain text file (UTF-8 recommended)."
        )

def extract_text_from_file(contents: bytes, filename: str) -> str:
    """Извлекает текст из файлов разных форматов (txt, doc, docx, pdf, odt)."""
    file_extension = os.path.splitext(filename.lower())[1]

    try:
        if file_extension in ['.txt', '.text']:
            return decode_file_contents(contents)

        elif file_extension in ['.doc', '.docx']:
            # Для DOCX используем python-docx напрямую из байтов
            if file_extension == '.docx':
                docx_obj = docx.Document(BytesIO(contents))
                text_parts = []
                for paragraph in docx_obj.paragraphs:
                    text_parts.append(paragraph.text)
                return '\n'.join(text_parts)
            else:  # .doc (старый формат)
                raise HTTPException(
                    status_code=400,
                    detail="DOC format is not fully supported. Please convert to DOCX or TXT."
                )

        elif file_extension == '.pdf':
            pdf_reader = PyPDF2.PdfReader(BytesIO(contents))
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            return '\n'.join(text_parts)

        elif file_extension == '.odt':
            doc = Document(BytesIO(contents))
            body = doc.body
            text_parts = []
            # Извлекаем текст из параграфов
            for paragraph in body.get_elements("text:p"):
                text = paragraph.text
                if text and text.strip():
                    text_parts.append(text.strip())

            # Извлекаем текст из заголовков
            for heading in body.get_elements("text:h"):
                text = heading.text
                if text and text.strip():
                    text_parts.append(text.strip())

            return '\n'.join(text_parts)

        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format: {file_extension}. Supported formats: txt, docx, pdf, odt."
            )

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file {filename}: {str(e)}"
        )

class SentimentSGR(BaseModel):
    sentiment_sgr: Literal["positive", "negative", "neutral"]
    confidence_sgr: float

class SummarySGR(BaseModel):
    summary_sgr: str
    original_text_length_sgr: int
    summary_length_sgr: int

def extract_sgr_json(raw_response: str) -> dict | None:
    """
    Извлекает JSON‑блок SGR из конца ответа модели.
    Возвращает словарь с данными SGR или None, если JSON не найден.
    """
    # Ищем JSON в конце строки (от { до })
    match = re.search(r'\{.*\}$', raw_response, re.DOTALL)
    if match:
        json_str = match.group(0)
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            pass
    return None


# База для анализа тональности
sentiment_documents = [
    {"text": "Отличный сервис, очень доволен!", "metadata": {"type": "sentiment", "label": "positive"}},
    {"text": "Ужасно, никогда больше не воспользуюсь", "metadata": {"type": "sentiment", "label": "negative"}},
    {"text": "Потрясающе! Сервис превзошёл все ожидания — быстрая доставка, вежливый персонал, качество на высоте. Обязательно буду рекомендовать друзьям!",
    "metadata": {"type": "sentiment", "label": "positive"}},
    {"text": "Крайне разочарован. Заказал товар неделю назад, до сих пор нет никакой информации. Поддержка не отвечает, деньги потрачены впустую. Никогда больше не буду пользоваться этой компанией.",
    "metadata": {"type": "sentiment", "label": "negative"}},
    {"text": "Получил заказ сегодня. Товар соответствует описанию, упаковка целая. Никаких проблем не возникло, но и особых восторгов тоже нет — просто выполнил свою функцию.",
    "metadata": {"type": "sentiment", "label": "neutral"}},
    {"text": "Огромное спасибо команде за отличную работу! Всё сделали быстро, чётко и с душой. Очень приятно иметь дело с профессионалами. Буду обращаться ещё!",
    "metadata": {"type": "sentiment", "label": "positive"}},
    {"text": "Ужасное обслуживание. Оператор грубил, решение проблемы затянулось на несколько дней. Качество товара тоже оставляет желать лучшего — видно, что сэкономили на материалах.",
    "metadata": {"type": "sentiment", "label": "negative"}},
    {"text": "Amazing experience! The product arrived earlier than expected, and it's even better than I imagined. Great value for money, highly recommend!",
    "metadata": {"type": "sentiment", "label": "positive"}},
    {"text": "Terrible service. Waited 3 hours on hold, then the agent couldn't help me. Product arrived damaged, and the company refuses to replace it. Very disappointed.",
    "metadata": {"type": "sentiment", "label": "negative"}},
    {"text": "Received the item as described. It works as expected, no issues so far. Nothing extraordinary, but it does the job. Average experience overall.",
    "metadata": {"type": "sentiment", "label": "neutral"}},
    {"text": "I'm very impressed with the quality and attention to detail. The team went above and beyond to meet my requirements. Will definitely order again!",
    "metadata": {"type": "sentiment", "label": "positive"}},
    {"text": "Waste of money. The product broke after one week of use. Customer service was unhelpful and rude. I regret this purchase and won't recommend this brand.",
    "metadata": {"type": "sentiment", "label": "negative"}},
    {"text": "The service is acceptable. It took longer than promised, but eventually I got what I ordered. Nothing to complain about, but also nothing to rave about.",
    "metadata": {"type": "sentiment", "label": "neutral"}}
]

# База для суммирования
summarize_documents = [
    {"text": "Вчера в Москве прошёл масштабный технологический форум, собравший более 5 000 специалистов из разных стран. Ключевыми темами мероприятия стали искусственный интеллект и кибербезопасность. Эксперты обсудили перспективы развития ИИ на ближайшие 5 лет и поделились лучшими практиками защиты данных. В рамках форума состоялось подписание нескольких важных соглашений между ведущими технологическими компаниями.",
    "summary": "В Москве прошёл технологический форум с участием 5 000 специалистов. Основные темы — ИИ и кибербезопасность. Обсуждены перспективы развития ИИ на 5 лет, подписаны межкорпоративные соглашения.",
    "metadata": {"type": "summarization", "domain": "news", "length_original": 342, "length_summary": 168}},
    {"text": "Купил этот ноутбук 3 месяца назад для работы и остался очень доволен. Процессор Intel i7 справляется с многозадачностью без проблем, 16 ГБ оперативной памяти хватает для всех задач. Экран яркий, цвета точные, что важно для дизайнера. Единственный минус — батарея держит всего 4 часа при активной работе. В целом, отличное соотношение цены и качества, рекомендую коллегам.",
    "summary": "Ноутбук с процессором Intel i7 и 16 ГБ ОЗУ хорошо подходит для работы, особенно для дизайна (точный цветопередача). Минус — батарея на 4 часа. Хорошее соотношение цены и качества.",
    "metadata": {"type": "summarization", "domain": "reviews", "length_original": 298, "length_summary": 158}},
    {"text": "Чтобы создать эффективное резюме, следуйте этим шагам: 1) Укажите ваше имя и контактную информацию в верхней части документа. 2) Кратко опишите ваш профессиональный опыт (не более 3–5 предложений). 3) Перечислите ключевые навыки, релевантные для желаемой должности. 4) Укажите образование и любые сертификаты. 5) При необходимости добавьте раздел «Достижения» с 2–3 пунктами. Старайтесь уместить всё на одной странице.",
    "summary": "Для создания резюме: укажите имя и контакты, кратко опишите опыт (3–5 предложений), перечислите ключевые навыки, образование и сертификаты. При необходимости добавьте 2–3 достижения. Уместите на одной странице.",
    "metadata": {"type": "summarization", "domain": "instructions", "length_original": 356, "length_summary": 182}},
    {"text": "A major technology conference was held in San Francisco last week, attracting over 7 000 attendees from 50 countries. Keynote speakers discussed advancements in renewable energy and electric vehicles. The event featured product launches from leading tech companies, including a new solar panel with 25 % higher efficiency. Several partnerships were announced between automotive and energy firms.",
    "summary": "San Francisco hosted a tech conference with 7 000+ attendees. Focus on renewable energy and EVs. New high‑efficiency solar panel launched; multiple automotive‑energy partnerships announced.",
    "metadata": {"type": "summarization", "domain": "news", "length_original": 289, "length_summary": 147}},
    {"text": "I've been using this wireless earbuds for two months now, and they've exceeded my expectations. The sound quality is excellent, with clear highs and deep bass. Battery life lasts 6 hours on a single charge, and the case provides three additional full charges. The only drawback is the touch controls are sometimes unresponsive. Overall, a great purchase for the price point.",
    "summary": "Wireless earbuds offer excellent sound (clear highs, deep bass) and 6‑hour battery (plus 3 charges from case). Minor issue: unresponsive touch controls. Good value overall.",
    "metadata": {"type": "summarization", "domain": "reviews", "length_original": 278, "length_summary": 139}},
    {"text": "To create an effective summary, follow these steps: 1) Read the original text thoroughly to understand the main points. 2) Identify the key ideas and supporting details. 3) Remove examples, repetitions, and minor details. 4) Combine related ideas into concise sentences. 5) Ensure the summary is 3–5 sentences long and preserves the original meaning. 6) Proofread for clarity and coherence.",
    "summary": "Create a summary by: reading thoroughly, identifying key ideas, removing extras, combining related points, keeping 3–5 sentences, and proofreading. Preserve original meaning.",
    "metadata": {"type": "summarization", "domain": "instructions", "length_original": 267, "length_summary": 143}},
    {"text": "Современные методы радиоволновой диагностики неоднородностей в ионосфере: теоретические основы и экспериментальные результаты\n\nИоносфера Земли представляет собой сложную плазменную среду, характеризующуюся высокой пространственно‑временной изменчивостью параметров. В последние десятилетия особый интерес вызывает изучение мелкомасштабных неоднородностей электронной концентрации, которые оказывают существенное влияние на распространение радиоволн различных диапазонов.\n\nТеоретическая модель, описывающая взаимодействие радиоволн с ионосферными неоднородностями, основывается на решении уравнения переноса излучения в случайно‑неоднородной среде. При этом учитываются следующие ключевые факторы:\n1) анизотропия флуктуаций электронной концентрации;\n2) зависимость коэффициента рассеяния от частоты радиоволны;\n3) эффекты многократного рассеяния;\n4) влияние геомагнитного поля на распространение волн.\n\nДля экспериментальной верификации теоретических предсказаний был проведён цикл измерений с использованием сети ионозондов вертикального зондирования и системы некогерентного рассеяния (ИСЗ «Ионосфера‑М»). Эксперименты проводились в период высокой солнечной активности (F10.7 > 180), что обеспечивало наличие интенсивных неоднородностей.\n\nОсновные результаты измерений:\n* зарегистрированы неоднородности с масштабами от 100 м до 50 км;\n* выявлена корреляция между интенсивностью неоднородностей и геомагнитной активностью (коэффициент корреляции r = 0.87);\n* обнаружено аномальное усиление флуктуаций фазы на частотах 10–30 МГц;\n* измерены вертикальные скорости перемещения неоднородностей (диапазон 50–200 м/с).\n\nОсобое внимание уделено анализу эффектов замирания сигнала (fading) при спутниковой связи. На основе полученных данных разработана методика коррекции ошибок, учитывающая:\n* статистические характеристики замираний;\n* пространственную структуру неоднородностей;\n* частотную зависимость коэффициента ослабления.\n\nПрактическая значимость результатов заключается в возможности:\n* повышения точности систем спутниковой навигации (GPS/ГЛОНАСС) на 15–20 % в условиях возмущённой ионосферы;\n* улучшения качества коротковолновой связи в полярных регионах;\n* оптимизации работы систем загоризонтной радиолокации.\n\nДальнейшие исследования планируется направить на изучение динамики неоднородностей в условиях различных геомагнитных бурь и разработку адаптивных алгоритмов компенсации ионосферных искажений.",
    "summary": "Исследование посвящено диагностике неоднородностей в ионосфере. Разработана теоретическая модель взаимодействия радиоволн с учётом анизотропии флуктуаций, частоты волны, многократного рассеяния и геомагнитного поля. Эксперименты с ионозондами и системой некогерентного рассеяния выявили неоднородности (100 м–50 км), их связь с геомагнитной активностью (r = 0.87), аномальное усиление флуктуаций фазы (10–30 МГц) и скорости перемещения (50–200 м/с). Разработана методика коррекции ошибок для спутниковой связи. Практическая польза: повышение точности GPS/ГЛОНАСС на 15–20 %, улучшение КВ‑связи в полярных регионах и оптимизация загоризонтной радиолокации.",
    "metadata": {
        "type": "summarization",
        "domain": "radio_physics",
        "topic": "ionospheric_irregularities",
        "length_original": 1384,
        "length_summary": 498,
        "key_elements_preserved": [
            "масштабы неоднородностей (100 м–50 км)",
            "коэффициент корреляции (r = 0.87)",
            "диапазон частот (10–30 МГц)",
            "скорости перемещения (50–200 м/с)",
            "повышение точности GPS/ГЛОНАСС (15–20 %)"
        ]}}
]

# Преобразуем словари в строки (если нужно)
if isinstance(sentiment_documents[0], dict):
    # Извлекаем только текст из поля "text"
    texts_sentiment_documents = [doc["text"] for doc in sentiment_documents]
else:
    # Оставляем как есть (если уже строки)
    texts_sentiment_documents = sentiment_documents

# Преобразуем словари в строки (если нужно)
if isinstance(summarize_documents[0], dict):
    # Извлекаем только текст из поля "text"
    texts_summarize_documents = [doc["text"] for doc in summarize_documents]
else:
    # Оставляем как есть (если уже строки)
    texts_summarize_documents = summarize_documents
    
# два отдельных ретривера
sentiment_retriever = HybridRetriever()
summarize_retriever = HybridRetriever()

# Загружаем соответствующие документы
sentiment_retriever.add_documents(texts_sentiment_documents)
summarize_retriever.add_documents(texts_summarize_documents)

app = FastAPI()


class TextRequest(BaseModel):
    text: str

@app.post("/analyze")
async def analyze_sentiment(
    text: str = Form(None),
    file: UploadFile = File(None)
):
    # Получаем текст из запроса или файла
    if file:
        contents = await file.read()
        text_content = extract_text_from_file(contents, file.filename)
    elif text is not None:
        text_content = text
    else:
        raise HTTPException(
            status_code=400,
            detail="Either text in request body or a file must be provided"
        )

    # Используем ретривер для поиска релевантных примеров
    relevant_examples = sentiment_retriever.retrieve(text_content, k=2)

    # Формируем промт с примерами из ретривера
    examples_section = ""
    if relevant_examples:
        examples_section = "\n## RELEVANT EXAMPLES FROM KNOWLEDGE BASE\n"
        for i, example in enumerate(relevant_examples, 1):
            examples_section += f"Example {i}:\n{example['text']}\nRelevance score: {example['score']:.2f}\n"

    prompt = f"""
## ROLE
You are a professional sentiment analysis expert. Your task is to accurately determine the emotional tone of text and assess confidence in your judgment.

## GOAL
Determine the sentiment of the text (positive, negative, neutral) and provide a confidence score (0.0 to 1.0).

## TASKS
1. Read the input text carefully.
2. Identify the overall emotional tone.
3. Evaluate confidence in the classification.
4. Provide output in the strictly defined format.

## CONTEXT
The text may be in English or Russian, contain informal language, slang, emojis, and typos. Analyze the overall mood, not individual words.

## INPUT DATA
Text for analysis:
---
{text_content}
---

## ANALYSIS INSTRUCTIONS
1. Positive sentiment: expresses joy, approval, gratitude, enthusiasm.
2. Negative sentiment: contains criticism, dissatisfaction, disappointment, anger.
3. Neutral sentiment: states facts without emotional coloring, contains technical information.

## OUTPUT FORMAT
Strictly follow this format:
- First line: one word — sentiment (`positive`, `negative`, or `neutral`).
- Second line: floating‑point number — confidence (from `0.0` to `1.0`, one decimal place).
- Third: a blank line.
- Fourth: JSON object with the following structure:
{{"sentiment_sgr": "positive|negative|neutral", "confidence_sgr": confidence_value}}

RULES FOR JSON:
- The JSON must be the LAST thing in your response.
- Do not include any text after the JSON object.
- The JSON must be valid and parsable.
- Replace `positive|negative|neutral` with the actual sentiment label.
- Replace `confidence_value` with the actual confidence score (e.g., 0.8).
- Ensure the JSON is on a separate line.

## EXAMPLES
Example 1:
Text: "Excellent service, very satisfied!"
Response:
positive
0.9

{{"sentiment_sgr": "positive", "confidence_sgr": 0.9}}

Example 2:
Text: "Nothing special, just an ordinary day"
Response:
neutral
0.7

{{"sentiment_sgr": "neutral", "confidence_sgr": 0.7}}

Example 3:
Text: "Terrible, I'll never use it again"
Response:
negative
0.8

{{"sentiment_sgr": "negative", "confidence_sgr": 0.8}}
"""

    
    try:
        response = ollama.chat(
            model='gemma3:12b',
            messages=[{'role': 'user', 'content': prompt}],
            options={'temperature': 0},
        )
        raw_output = response['message']['content'].strip()

        # Извлекаем SGR‑данные
        sgr_data = extract_sgr_json(raw_output)

        if sgr_data:
            # Валидируем через Pydantic
            sgr_result = SentimentSGR(**sgr_data)
            return {
                "sentiment": sgr_result.sentiment_sgr,
                "confidence": sgr_result.confidence_sgr,
                "sgr_raw": raw_output  # опционально: полный ответ с JSON
            }
        else:
            # Fallback: парсим старый формат (если SGR не сработал)
            lines = raw_output.split('\n')
            if len(lines) < 2:
                raise HTTPException(status_code=500, detail="Invalid LLM response format")

            sentiment = lines[0].lower()
            try:
                confidence = float(lines[1])
            except ValueError:
                raise HTTPException(status_code=500, detail="Invalid confidence value")

            if sentiment not in ["positive", "negative", "neutral"]:
                raise HTTPException(status_code=500, detail="Invalid sentiment label")
            if not (0 <= confidence <= 1):
                raise HTTPException(status_code=500, detail="Confidence must be between 0 and 1")

            return {"sentiment": sentiment, "confidence": confidence}

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error communicating with Ollama: {str(e)}"
        )

@app.post("/summarize")
async def summarize_text(
    text: str = Form(None),
    file: UploadFile = File(None)
):
    """
    Endpoint for text summarization. Accepts either text in request body or a text file.
    """
    # Получаем текст из запроса или файла
    if file:
        contents = await file.read()
        text_content = extract_text_from_file(contents, file.filename)
    elif text is not None:
        text_content = text
    else:
        raise HTTPException(
            status_code=400,
            detail="Either text in request body or a file must be provided"
        )

    # Ищем релевантные документы для резюмирования
    relevant_docs = summarize_retriever.retrieve(text_content, k=1)
    

    context_section = ""
    if relevant_docs:
        context_section = f"\n## CONTEXTUAL GUIDANCE\nUse these guidelines for summarization:\n{relevant_docs[0]['text']}"

    # Create summarization prompt
    prompt = f"""
## ROLE
You are an experienced editor and text analyst. Your task is to create concise, informative summaries of any texts while preserving key information.

## GOAL
Create a structured summary of 3–5 sentences that conveys main ideas and important details without extra information.

## TASKS
1. Read and analyze the full text.
2. Extract key facts, main ideas, and critical conclusions.
3. Eliminate repetitions and minor details.
4. Formulate 3–5 coherent sentences reflecting the essence of the text.
5. Ensure the summary reads as a standalone piece.

## CONTEXT
The text may contain technical information, news, articles, reviews, or informal communication. The summary must be understandable without reading the original.

## INPUT DATA
Original text:
---
{text_content}
---

## SUMMARY CREATION INSTRUCTIONS
1. Focus on the main topic and key conclusions.
2. Remove repetitions and redundant details.
3. Preserve critical numbers, dates, and names if essential.
4. Use simple, clear language.
5. Maintain logical coherence between sentences.

## QUALITY CRITERIA
- Length: exactly 3–5 sentences.
- Completeness: all key ideas preserved.
- Conciseness: examples and minor details excluded.
- Readability: text is easily comprehensible.

## OUTPUT FORMAT
1. First, provide the final summary — 3–5 consecutive sentences. Do not include any headings, comments, or formatting.
2. Then, add a blank line.
3. Finally, add the following JSON object:
{{"summary_sgr": "your_summary_here", "original_text_length_sgr": original_length, "summary_length_sgr": summary_length}}

RULES FOR JSON:
- The JSON must be the LAST thing in your response.
- Do not include any text after the JSON object.
- Replace `your_summary_here` with the actual summary text.
- Replace `original_length` with the character count of the original text.
- Replace `summary_length` with the character count of your summary.
- Ensure the JSON is valid and parsable.

## EXAMPLES
Example 1:
Original text: "Yesterday, a major technology forum was held in Moscow. It was attended by over 5 000 professionals from different countries. Key topics included artificial intelligence and cybersecurity. Experts discussed AI development prospects for the next 5 years and shared best practices for data protection. The event concluded with the signing of several important agreements between companies."

Summary:
A technology forum took place in Moscow with over 5 000 professionals in attendance. The main topics were artificial intelligence and cybersecurity. Experts reviewed AI development prospects for the next 5 years and discussed data protection methods. The forum resulted in important inter‑company agreements being signed.

{{"summary_sgr": "A technology forum took place in Moscow with over 5 000 professionals in attendance. The main topics were artificial intelligence and cybersecurity. Experts reviewed AI development prospects for the next 5 years and discussed data protection methods. The forum resulted in important inter‑company agreements being signed.", "original_text_length_sgr": 346, "summary_length_sgr": 258}}
"""


    try:
        response = ollama.chat(
            model='gemma3:12b',
            messages=[{'role': 'user', 'content': prompt}],
            options={'temperature': 0.3, 'num_predict': 300}
        )
        raw_output = response['message']['content'].strip()

        # Извлекаем SGR‑данные
        sgr_data = extract_sgr_json(raw_output)

        if sgr_data:
            # Валидируем через Pydantic
            sgr_result = SummarySGR(**sgr_data)
            return {
                "summary": sgr_result.summary_sgr,
                "original_text_length": sgr_result.original_text_length_sgr,
                "summary_length": sgr_result.summary_length_sgr,
                "sgr_raw": raw_output  # опционально: полный ответ с JSON
            }
        else:
            # Fallback: используем старый способ
            summary = raw_output
            if not summary or len(summary) < 10:
                raise HTTPException(
                    status_code=500,
                    detail="Generated summary is too short or empty"
                )

            return {
                "original_text_length": len(text_content),
                "summary_length": len(summary),
                "summary": summary
            }

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error during summarization: {str(e)}"
        )
