import pandas as pd
import requests
import os
from pathlib import Path
from typing import List, Dict, Any
from io import BytesIO
import docx
import PyPDF2
from odfdo import Document as OdfDocument

class MultiFileAPITester:
    def __init__(self, api_base_url: str = "http://localhost:8000"):
        self.api_base_url = api_base_url
        self.session = requests.Session()

    def extract_text_from_file(self, file_path: Path) -> str:
        """Извлекает текст из файла любого поддерживаемого формата."""
        file_extension = file_path.suffix.lower()

        with open(file_path, 'rb') as f:
            contents = f.read()

        if file_extension in ['.txt', '.text']:
            return self.decode_file_contents(contents)

        elif file_extension == '.docx':
            docx_obj = docx.Document(BytesIO(contents))
            text_parts = []
            # Извлекаем текст из параграфов
            for paragraph in docx_obj.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # Извлекаем текст из таблиц
            for table in docx_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return '\n'.join(text_parts)

        elif file_extension == '.pdf':
            pdf_reader = PyPDF2.PdfReader(BytesIO(contents))
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            return '\n'.join(text_parts)

        elif file_extension == '.odt':
            doc = OdfDocument(BytesIO(contents))
            body = doc.body
            text_parts = []
            for paragraph in body.get_elements("text:p"):
                text = paragraph.text
                if text and text.strip():
                    text_parts.append(text.strip())
            return '\n'.join(text_parts)

        else:
            raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")

    def decode_file_contents(self, contents: bytes) -> str:
        """Декодирование текстовых файлов с разными кодировками."""
        encodings_to_try = ['utf-8', 'cp1251', 'koi8-r', 'iso-8859-5', 'windows-1252', 'ascii']

        for encoding in encodings_to_try:
            try:
                return contents.decode(encoding)
            except UnicodeDecodeError:
                continue

        try:
            return contents.decode('utf-8', errors='replace')
        except Exception:
            raise ValueError("Не удалось декодировать файл")

    def load_dataset_from_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Загружает все поддерживаемые файлы из директории."""
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Директория не найдена: {directory_path}")

        supported_extensions = {'.txt', '.docx', '.pdf', '.odt'}
        files = [f for f in directory.iterdir() if f.is_file() and f.suffix.lower() in supported_extensions]

        dataset = []
        for file_path in files:
            try:
                text_content = self.extract_text_from_file(file_path)
                dataset.append({
                    'filename': file_path.name,
            'text': text_content,
            'file_path': str(file_path)
                })
            except Exception as e:
                print(f"Ошибка обработки файла {file_path}: {e}")

        return dataset

    def send_request(self, text: str, endpoint: str) -> Dict[str, Any]:
        """Отправляет запрос к указанному эндпоинту API."""
        url = f"{self.api_base_url}/{endpoint}"
        files = {'file': ('input.txt', text, 'text/plain')}

        try:
            response = self.session.post(url, files=files, timeout=180)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            return {
                'error': str(e),
                'result': f"Ошибка API: {str(e)}",
                'request_text': text
            }

    def process_dataset(
        self,
        dataset_directory: str,
        task_type: str = 'analyze'
    ) -> List[Dict[str, Any]]:
        """Обрабатывает все файлы в директории и собирает результаты."""
        print(f"Загрузка файлов из директории: {dataset_directory}")
        data = self.load_dataset_from_directory(dataset_directory)
        print(f"Загружено {len(data)} файлов")

        results = []
        for i, item in enumerate(data, 1):
            print(f"Обработка файла {i}/{len(data)}: {item['filename']}")

            text = item['text']
            if not text:
                results.append({
                    'запрос': f"Файл: {item['filename']}",
            'результат': 'Ошибка: пустой файл',
            'original_text': ''
                })
                continue

            result = self.send_request(text, task_type)

            if task_type == 'analyze':
                formatted_result = (
                    f"Sentiment: {result.get('sentiment', 'N/A')}, "
            f"Confidence: {result.get('confidence', 'N/A')}"
                )
            else:  # summarize
                formatted_result = result.get('summary', 'N/A')

            results.append({
                'запрос': f"Файл: {item['filename']} (текст: {text[:150]}...)",
            'результат': formatted_result,
            'original_text': text,
            'filename': item['filename']
            })

        return results

    def save_results(self, results: List[Dict[str, Any]], output_path: str):
        """Сохраняет результаты в CSV‑файл."""
        table_data = [
            {
                'запрос': item['запрос'],
            'результат': item['результат'],
            'имя_файла': item['filename']
            }
            for item in results
        ]

        df = pd.DataFrame(table_data)
        output_path = Path(output_path)
        df.to_csv(output_path, index=False, encoding='utf-8-sig')
        print(f"Результаты сохранены в: {output_path}")

def main():
    # Настройки
    API_URL = "http://localhost:8000"
    DATASET_DIRECTORY = "dataset_files"  # Директория с файлами
    OUTPUT_PATH = "results_from_files.csv"
    TASK_TYPE = "analyze"  # "analyze" или "summarize"

    tester = MultiFileAPITester(api_base_url=API_URL)

    try:
        results = tester.process_dataset(DATASET_DIRECTORY, TASK_TYPE)
        tester.save_results(results, OUTPUT_PATH)
    except Exception as e:
        print(f"Произошла ошибка: {e}")

if __name__ == "__main__":
    main()
